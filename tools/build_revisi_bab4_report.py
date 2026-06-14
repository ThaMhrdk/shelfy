from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\Program Files\xampp\htdocs\laravel\shelfy")
SOURCE = Path(r"C:\Users\THA\Downloads\Laporan Basdat\Laporan TuBes - Shelfy (Tyas).docx")
OUTPUT = ROOT / "docs" / "Laporan TUBES SHELFY - Revisi BAB IV.docx"
ASSETS = ROOT / "docs" / "assets-shelfy-report"

BLUE = "1F4E79"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F2F2"
PLACEHOLDER_FILL = "FFF2CC"
PLACEHOLDER_BORDER = "BF9000"


def element_text(element):
    return "".join(node.text or "" for node in element.iter(qn("w:t"))).strip()


def remove_from_chapter_four(document):
    body = document._element.body
    removing = False

    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue

        text = " ".join(element_text(child).split()).upper()
        if not removing and text.startswith("BAB IV IMPLEMENTASI"):
            removing = True

        if removing:
            body.remove(child)

    if not removing:
        raise RuntimeError("BAB IV IMPLEMENTASI tidak ditemukan pada dokumen sumber.")


def clean_base_chapters(document):
    remove_markers = (
        "ISILAH TITIK-TITIK SESEUAI DENGAN YANG KALIAN BUAT",
        "JELASKAN MASING-MASING BAGIAN SECARA SINGKAT",
    )

    for paragraph in list(document.paragraphs):
        text = " ".join(paragraph.text.split())
        upper = text.upper()

        if any(marker in upper for marker in remove_markers):
            element = paragraph._element
            element.getparent().remove(element)
            continue

        if upper.startswith("3.3.1 STRUKTUR DATABASE"):
            paragraph.text = "3.3.1 Struktur Database"
            paragraph.style = document.styles["Heading 2"]
            continue

        if text == "3.3.3 Relasi Antar Kota":
            paragraph.text = "3.3.3 Relasi Antar Data"
            continue

        if text == "Perancangan Stuktur Folder":
            paragraph.text = "Perancangan Struktur Folder"


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_border(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for edge in ("top", "left", "bottom", "right"):
        tag = qn(f"w:{edge}")
        border = borders.find(tag)
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "10")
        border.set(qn("w:color"), color)


def format_normal(paragraph):
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(11)


def add_body(document, text):
    paragraph = document.add_paragraph(text)
    format_normal(paragraph)
    return paragraph


def add_bullet(document, text):
    paragraph = document.add_paragraph(text, style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.1
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(11)
    return paragraph


def add_heading(document, text, level=1):
    paragraph = document.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(10 if level > 1 else 14)
    paragraph.paragraph_format.space_after = Pt(5)
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.bold = True
        run.font.color.rgb = RGBColor(31, 78, 121)
        run.font.size = Pt({1: 15, 2: 13, 3: 11.5}.get(level, 11))
    return paragraph


def add_chapter(document, roman, title):
    document.add_page_break()
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(f"BAB {roman}")
    run.font.name = "Arial"
    run.font.size = Pt(16)
    run.font.bold = True
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(18)
    run = paragraph.add_run(title.upper())
    run.font.name = "Arial"
    run.font.size = Pt(16)
    run.font.bold = True


def add_code_placeholder(document, title, paths, capture_hint):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.25)
    cell = table.cell(0, 0)
    cell.width = Inches(6.25)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, PLACEHOLDER_FILL)
    set_cell_border(cell, PLACEHOLDER_BORDER)
    prevent_row_split(table.rows[0])

    heading = cell.paragraphs[0]
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_after = Pt(6)
    run = heading.add_run("TEMPAT SCREENSHOT VISUAL STUDIO CODE")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(127, 96, 0)

    label = cell.add_paragraph()
    label.paragraph_format.space_after = Pt(4)
    run = label.add_run(f"Bagian yang difoto: {title}")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(10)

    for path in paths:
        paragraph = cell.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.15)
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(f"Path: {ROOT / path}")
        run.font.name = "Consolas"
        run.font.size = Pt(9)

    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(f"Petunjuk foto: {capture_hint}")
    run.italic = True
    run.font.name = "Arial"
    run.font.size = Pt(9.5)

    caption = document.add_paragraph(f"Placeholder screenshot kode: {title}")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(8)
    for run in caption.runs:
        run.italic = True
        run.font.name = "Arial"
        run.font.size = Pt(9)


def add_image(document, filename, caption, width=6.25):
    path = ASSETS / filename
    if not path.exists():
        raise FileNotFoundError(path)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(path), width=Inches(width))

    caption_paragraph = document.add_paragraph(caption)
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_after = Pt(8)
    for run in caption_paragraph.runs:
        run.italic = True
        run.font.name = "Arial"
        run.font.size = Pt(9)


def add_table(document, headers, rows, widths=None):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False

    if widths is None:
        widths = [6.25 / len(headers)] * len(headers)

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.width = Inches(widths[index])
        set_cell_shading(cell, LIGHT_BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(9.5)
    set_repeat_header(table.rows[0])
    prevent_row_split(table.rows[0])

    for values in rows:
        row = table.add_row()
        prevent_row_split(row)
        for index, value in enumerate(values):
            cell = row.cells[index]
            cell.width = Inches(widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.space_before = Pt(2)
            run = paragraph.add_run(str(value))
            run.font.name = "Arial"
            run.font.size = Pt(9)

    document.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def build():
    document = Document(SOURCE)
    clean_base_chapters(document)
    remove_from_chapter_four(document)

    add_chapter(document, "IV", "Implementasi")
    add_heading(document, "Implementasi", 1)
    add_body(
        document,
        "Implementasi merupakan tahap penerapan hasil analisis dan perancangan menjadi aplikasi yang dapat digunakan. "
        "SHELFY dibangun menggunakan framework Laravel dengan pola Model-View-Controller (MVC), Laravel Breeze untuk "
        "autentikasi, Blade untuk antarmuka, MongoDB sebagai database NoSQL, dan DomPDF/fitur cetak browser untuk nota. "
        "Aplikasi memiliki tiga role, yaitu admin, pustakawan, dan mahasiswa.",
    )
    add_body(
        document,
        "Admin bertugas memantau sistem dan memiliki hak menghapus data buku. Pustakawan menjalankan kegiatan operasional "
        "perpustakaan seperti mengelola buku, mencatat pengambilan, memproses pengembalian, perpanjangan, serta konfirmasi "
        "pembayaran denda. Mahasiswa menggunakan katalog, keranjang, checkout, peminjaman saya, dan riwayat nota.",
    )

    add_heading(document, "Implementasi Backend", 1)
    add_body(
        document,
        "Backend SHELFY memanfaatkan route Laravel untuk menghubungkan permintaan pengguna dengan Controller. Controller "
        "melakukan validasi, mengambil atau memperbarui dokumen MongoDB melalui Model, kemudian mengirim data ke Blade view. "
        "Route utama berada pada routes/web.php dan route autentikasi Breeze berada pada routes/auth.php.",
    )

    add_heading(document, "Implementasi Login dan Register", 2)
    add_body(
        document,
        "Proses login ditangani oleh AuthenticatedSessionController. Method store memanggil proses autentikasi Breeze, "
        "meregenerasi session, kemudian mengarahkan pengguna ke dashboard sesuai role melalui Shelfy::homeRouteName. "
        "Middleware memastikan halaman petugas hanya dapat dibuka admin atau pustakawan, sedangkan halaman mahasiswa hanya "
        "dapat dibuka akun dengan role mahasiswa.",
    )
    add_body(
        document,
        "Proses register ditangani oleh RegisteredUserController. Sistem memvalidasi nama, NIM, program studi, kontak, email, "
        "dan password. Register membuat dokumen pada collection members dan users, menghubungkan keduanya melalui member_id "
        "dan user_id, lalu mengarahkan pengguna kembali ke halaman login. Pengguna tidak langsung masuk ke aplikasi setelah "
        "melakukan register.",
    )
    add_code_placeholder(
        document,
        "Login, register, dan pengalihan dashboard berdasarkan role",
        [
            Path("app/Http/Controllers/Auth/AuthenticatedSessionController.php"),
            Path("app/Http/Controllers/Auth/RegisteredUserController.php"),
            Path("app/Support/Shelfy.php"),
        ],
        "Buka method store pada kedua Controller dan method homeRouteName pada Shelfy.php. Foto bagian kode yang memperlihatkan validasi, pembuatan users/members, redirect ke login, dan redirect sesuai role.",
    )

    add_heading(document, "Implementasi Pengelolaan Buku", 2)
    add_body(
        document,
        "Pengelolaan buku dilakukan melalui BookController. Method index mengambil seluruh data buku lalu menerapkan filter "
        "judul, penulis, ISBN, kategori, dan status stok. Method store digunakan untuk menambah sekaligus memperbarui buku. "
        "Pada proses edit, stok tersedia dihitung kembali berdasarkan jumlah buku yang sedang dipinjam agar perubahan stok "
        "total tidak merusak data transaksi.",
    )
    add_body(
        document,
        "Cover buku divalidasi sebagai gambar dan disimpan pada disk public di folder storage/app/public/covers. Path cover "
        "disimpan pada collection books dan dapat diakses melalui public/storage setelah symbolic link dibuat. Penghapusan "
        "buku hanya dapat dilakukan admin dan ditolak apabila buku masih memiliki peminjaman aktif.",
    )
    add_code_placeholder(
        document,
        "CRUD buku, filter, perhitungan stok, dan upload cover",
        [Path("app/Http/Controllers/BookController.php"), Path("app/Models/Book.php")],
        "Foto method index, store, destroy, filterBooks, dan storeCover. Jika satu layar tidak cukup, gunakan dua screenshot dan letakkan berurutan pada kotak ini.",
    )

    add_heading(document, "Implementasi Keranjang dan Checkout Peminjaman", 2)
    add_body(
        document,
        "Mahasiswa memilih buku dari katalog dan mengisi tanggal peminjaman melalui popup. StudentController::addToCart "
        "memeriksa status anggota, ketersediaan stok, serta mencegah item ganda. Data sementara disimpan pada collection "
        "cart_items sehingga mahasiswa dapat memeriksa kembali pilihan sebelum checkout.",
    )
    add_body(
        document,
        "Saat checkout, sistem memeriksa seluruh item dan ketersediaan buku sekali lagi. Setiap item kemudian diubah menjadi "
        "dokumen pada collection loans dengan status awal menunggu_diambil. Stok tersedia dikurangi, jumlah peminjaman buku "
        "ditambah, dan item keranjang dihapus. Status ini memberi waktu kepada pustakawan untuk memastikan buku benar-benar "
        "telah diserahkan kepada mahasiswa.",
    )
    add_code_placeholder(
        document,
        "Masuk keranjang dan checkout menjadi peminjaman",
        [Path("app/Http/Controllers/StudentController.php"), Path("app/Models/CartItem.php"), Path("app/Models/Loan.php")],
        "Foto method addToCart dan checkout pada StudentController. Sertakan bagian pembuatan dokumen Loan dan perubahan stok buku.",
    )

    add_heading(document, "Implementasi Bukti Buku Sudah Diambil", 2)
    add_body(
        document,
        "Setelah checkout, data peminjaman muncul pada halaman petugas dengan status menunggu_diambil. Pustakawan menggunakan "
        "LoanController::pickup untuk mencatat tanggal pengambilan, keterangan bukti pengambilan, dan nama petugas. Status "
        "berubah menjadi dipinjam atau terlambat apabila jatuh tempo telah terlewati.",
    )
    add_code_placeholder(
        document,
        "Pencatatan bukti buku sudah diambil",
        [Path("app/Http/Controllers/LoanController.php")],
        "Foto method pickup dan refreshOverdue agar terlihat perubahan status menunggu_diambil menjadi dipinjam/terlambat.",
    )

    add_heading(document, "Implementasi Pengembalian dan Perpanjangan", 2)
    add_body(
        document,
        "Pengembalian diproses oleh ReturnController::store. Pustakawan wajib memasukkan tanggal kembali secara manual agar "
        "sesuai dengan kondisi nyata. Sistem menolak tanggal setelah hari ini atau tanggal sebelum tanggal pinjam. Setelah "
        "valid, Shelfy::lateFee menghitung hari keterlambatan dan denda, status pinjaman berubah menjadi dikembalikan, nomor "
        "nota dibuat, dan stok tersedia bertambah.",
    )
    add_body(
        document,
        "Perpanjangan diproses oleh ReturnController::extend. Pustakawan mengisi tanggal jatuh tempo baru dan alasan "
        "perpanjangan. Sistem memastikan pinjaman masih aktif, belum terlambat, serta tanggal baru melewati jatuh tempo lama. "
        "Riwayat perpanjangan disimpan sebagai array embedded document pada field riwayat_perpanjangan di collection loans.",
    )
    add_code_placeholder(
        document,
        "Pengembalian manual, kalkulasi denda, dan riwayat perpanjangan",
        [Path("app/Http/Controllers/ReturnController.php"), Path("app/Support/Shelfy.php")],
        "Foto method store dan extend pada ReturnController serta method lateFee pada Shelfy.php.",
    )

    add_heading(document, "Implementasi Nota dan Pembayaran Denda", 2)
    add_body(
        document,
        "Nota hanya dapat dibuka setelah peminjaman berstatus dikembalikan. ReceiptController::show memastikan mahasiswa hanya "
        "dapat melihat nota miliknya, sedangkan admin dan pustakawan dapat melihat nota operasional. Nota menampilkan nomor "
        "nota, identitas anggota, buku, tanggal transaksi, keterlambatan, denda, pembayaran, dan petugas.",
    )
    add_body(
        document,
        "Apabila terdapat denda, mahasiswa memilih metode pembayaran Transfer atau QRIS melalui method pay. Status pembayaran "
        "menjadi menunggu_konfirmasi. Admin atau pustakawan kemudian menjalankan method confirm untuk mengubah status menjadi "
        "lunas. Halaman nota dapat dicetak atau disimpan sebagai PDF melalui tampilan cetak browser.",
    )
    add_code_placeholder(
        document,
        "Nota, pembayaran denda, dan konfirmasi pustakawan",
        [Path("app/Http/Controllers/ReceiptController.php"), Path("resources/views/shelfy/receipts/show.blade.php")],
        "Foto method show, pay, dan confirm pada ReceiptController. Tambahkan screenshot Blade nota bila diperlukan.",
    )

    add_heading(document, "Implementasi Rekapitulasi Data", 2)
    add_body(
        document,
        "DashboardController dan RecapController melakukan rekapitulasi dari dokumen books dan loans. Operasi yang digunakan "
        "meliputi count untuk jumlah judul dan transaksi, sum untuk stok total dan stok tersedia, avg untuk rata-rata stok, "
        "greater than untuk menghitung buku dengan stok lebih dari lima, serta groupBy untuk rekap kategori dan buku populer. "
        "Hasil rekap ditampilkan pada dashboard admin dan pustakawan.",
    )
    add_code_placeholder(
        document,
        "Rekapitulasi SUM, AVG, greater than, groupBy, dan buku populer",
        [Path("app/Http/Controllers/DashboardController.php"), Path("app/Http/Controllers/RecapController.php")],
        "Foto method stats, categoryLoans, categoryStats, serta pemanggilan Shelfy::topBorrowed.",
    )

    add_heading(document, "Implementasi Frontend", 1)
    add_body(
        document,
        "Frontend SHELFY menggunakan Laravel Blade, CSS, dan JavaScript ringan untuk popup atau modal. Layout menampilkan "
        "sidebar sesuai role, topbar pencarian, tabel operasional untuk petugas, serta katalog berbentuk kartu untuk mahasiswa. "
        "Berikut tampilan utama yang telah diimplementasikan.",
    )

    add_heading(document, "Login dan Register", 2)
    add_image(document, "01-login.png", "Gambar 4.1 Halaman login SHELFY", 5.3)
    add_body(
        document,
        "Halaman login menjadi pintu masuk seluruh pengguna. Pengguna memasukkan email dan password, kemudian sistem "
        "mengarahkan akun ke bagian admin/pustakawan atau mahasiswa sesuai role.",
    )
    add_image(document, "02-register.png", "Gambar 4.2 Halaman register mahasiswa", 5.3)
    add_body(
        document,
        "Halaman register digunakan untuk pendaftaran mandiri mahasiswa. Setelah registrasi berhasil, pengguna diarahkan ke "
        "halaman login dan harus masuk menggunakan akun yang baru dibuat.",
    )

    add_heading(document, "Dashboard", 2)
    add_heading(document, "Dashboard Admin dan Pustakawan", 3)
    add_image(document, "03-admin-dashboard.png", "Gambar 4.3 Dashboard admin/pustakawan")
    add_body(
        document,
        "Dashboard petugas menampilkan total judul, jumlah stok tersedia, peminjaman aktif, keterlambatan, rekap buku per "
        "kategori, buku terbaru, serta informasi cepat transaksi. Data ini membantu petugas memantau kondisi perpustakaan.",
    )
    add_heading(document, "Dashboard Mahasiswa", 3)
    add_image(document, "10-student-dashboard.png", "Gambar 4.4 Dashboard mahasiswa")
    add_body(
        document,
        "Dashboard mahasiswa menampilkan ringkasan peminjaman aktif, keterlambatan, transaksi selesai, daftar peminjaman "
        "terbaru, buku tersedia, dan buku populer.",
    )

    add_heading(document, "Bagian Admin dan Pustakawan", 2)
    add_heading(document, "Pengelolaan Buku", 3)
    add_image(document, "04-admin-books.png", "Gambar 4.5 Halaman pengelolaan buku")
    add_body(
        document,
        "Halaman Buku menyediakan pencarian, filter kategori, filter stok, detail buku, tambah buku, edit melalui popup, "
        "upload cover, dan penghapusan khusus admin. Admin dan pustakawan dapat menyesuaikan stok sesuai buku fisik.",
    )
    add_heading(document, "Direktori Anggota", 3)
    add_image(document, "05-admin-members.png", "Gambar 4.6 Halaman direktori anggota")
    add_body(
        document,
        "Halaman Anggota menampilkan mahasiswa yang terdaftar beserta NIM, program studi, kontak, status, dan jumlah "
        "peminjaman aktif. Penambahan anggota dilakukan melalui register mahasiswa agar alur akun lebih wajar.",
    )
    add_heading(document, "Peminjaman dan Bukti Pengambilan", 3)
    add_image(document, "06-admin-loans.png", "Gambar 4.7 Halaman peminjaman dan bukti pengambilan")
    add_body(
        document,
        "Halaman Peminjaman digunakan petugas untuk memantau transaksi dari checkout mahasiswa dan mengonfirmasi bahwa buku "
        "telah diambil. Pencarian dan filter status memudahkan pemantauan transaksi.",
    )
    add_heading(document, "Pengembalian, Perpanjangan, dan Nota", 3)
    add_image(document, "07-pustakawan-returns.png", "Gambar 4.8 Halaman pengembalian dengan tanggal manual")
    add_image(document, "08-popup-perpanjang.png", "Gambar 4.9 Popup perpanjangan peminjaman")
    add_body(
        document,
        "Pustakawan memasukkan tanggal kembali manual ketika memproses pengembalian. Apabila mahasiswa membutuhkan tambahan "
        "waktu, pustakawan membuka popup Perpanjang dan mengisi jatuh tempo baru serta alasannya.",
    )
    add_image(document, "09-nota-pdf.png", "Gambar 4.10 Nota pengembalian yang dapat dicetak atau disimpan sebagai PDF")
    add_body(
        document,
        "Nota pengembalian menjadi bukti formal transaksi. Nota dapat dicetak menggunakan printer atau disimpan sebagai file "
        "PDF melalui menu cetak browser.",
    )

    add_heading(document, "Bagian Mahasiswa", 2)
    add_heading(document, "Katalog dan Detail Buku", 3)
    add_image(document, "11-student-catalog-5kolom.png", "Gambar 4.11 Katalog mahasiswa lima kolom")
    add_body(
        document,
        "Katalog menampilkan lima kartu buku dalam satu baris desktop. Mahasiswa dapat melakukan pencarian, filter kategori, "
        "filter stok, membuka detail, dan memasukkan buku yang tersedia ke keranjang.",
    )
    add_image(document, "15-student-book-detail.png", "Gambar 4.12 Halaman detail buku")
    add_body(
        document,
        "Halaman detail menampilkan cover, metadata, deskripsi, dan stok buku. Tombol masuk keranjang membuka form peminjaman "
        "tanpa mengharuskan mahasiswa berpindah ke form administrasi petugas.",
    )
    add_heading(document, "Keranjang dan Checkout", 3)
    add_image(document, "12-student-cart.png", "Gambar 4.13 Keranjang dan checkout peminjaman")
    add_body(
        document,
        "Keranjang berfungsi sebagai tahap pemeriksaan sebelum checkout. Mahasiswa dapat menghapus pilihan atau menyelesaikan "
        "checkout untuk membuat transaksi peminjaman.",
    )
    add_heading(document, "Peminjaman Saya dan Riwayat Nota", 3)
    add_image(document, "13-student-loans.png", "Gambar 4.14 Halaman Peminjaman Saya")
    add_body(
        document,
        "Peminjaman Saya menampilkan status transaksi aktif dan detail peminjaman. Mahasiswa dapat melihat apakah buku masih "
        "menunggu diambil, sedang dipinjam, terlambat, atau telah dikembalikan.",
    )
    add_image(document, "14-student-history.png", "Gambar 4.15 Halaman Riwayat dan Nota")
    add_body(
        document,
        "Riwayat dan Nota menampilkan transaksi yang telah selesai. Mahasiswa dapat membuka nota pengembalian dan melakukan "
        "pembayaran denda apabila terdapat keterlambatan.",
    )

    add_heading(document, "Koneksi MongoDB", 1)
    add_body(
        document,
        "Koneksi MongoDB dikonfigurasi melalui file .env dengan DB_CONNECTION=mongodb, host 127.0.0.1, port 27017, dan nama "
        "database shelfy_db. File config/database.php mendefinisikan driver mongodb dari package mongodb/laravel-mongodb. "
        "Setiap Model utama mewarisi MongoDB\\Laravel\\Eloquent\\Model, sedangkan User menggunakan MongoDB Authenticatable agar "
        "fitur login Laravel Breeze dapat membaca collection users.",
    )
    add_body(
        document,
        "Untuk memeriksa data melalui MongoDB Compass, koneksi lokal yang digunakan adalah mongodb://127.0.0.1:27017. Setelah "
        "terhubung, database shelfy_db menampilkan collection utama users, members, books, loans, dan cart_items. File cover "
        "buku serta foto profil tidak disimpan langsung sebagai binary di MongoDB; database hanya menyimpan path file pada "
        "storage Laravel.",
    )
    add_code_placeholder(
        document,
        "Konfigurasi koneksi MongoDB dan Model MongoDB",
        [Path(".env"), Path("config/database.php"), Path("app/Models/User.php"), Path("app/Models/Book.php")],
        "Foto bagian DB_CONNECTION sampai DB_DATABASE pada .env, blok koneksi mongodb pada config/database.php, dan deklarasi class Model yang menggunakan package MongoDB.",
    )

    document.add_page_break()
    add_heading(document, "Pengujian Aplikasi", 1)
    add_body(
        document,
        "Pengujian dilakukan dengan membuka aplikasi pada browser lokal, menjalankan route Laravel, dan memeriksa perubahan "
        "dokumen melalui MongoDB Compass. Tabel berikut merangkum pengujian fungsi utama.",
    )
    add_table(
        document,
        ["No.", "Skenario Pengujian", "Hasil yang Diharapkan", "Status"],
        [
            ["1", "Register mahasiswa", "Dokumen users dan members dibuat, lalu diarahkan ke login.", "Berhasil"],
            ["2", "Login tiga role", "Admin/pustakawan masuk dashboard petugas; mahasiswa masuk dashboard mahasiswa.", "Berhasil"],
            ["3", "CRUD, filter, dan upload cover buku", "Data books berubah dan cover dapat ditampilkan melalui storage.", "Berhasil"],
            ["4", "Tambah buku ke keranjang dan checkout", "cart_items berubah menjadi loans berstatus menunggu_diambil.", "Berhasil"],
            ["5", "Konfirmasi buku diambil", "Status berubah menjadi dipinjam dan bukti pengambilan tersimpan.", "Berhasil"],
            ["6", "Perpanjangan", "Jatuh tempo baru dan riwayat perpanjangan tersimpan.", "Berhasil"],
            ["7", "Pengembalian tanggal manual", "Status menjadi dikembalikan, stok bertambah, dan denda dihitung.", "Berhasil"],
            ["8", "Nota dan pembayaran denda", "Nota tampil, pembayaran dicatat, dan pustakawan dapat konfirmasi lunas.", "Berhasil"],
            ["9", "Rekapitulasi", "SUM, AVG, greater than, kategori, dan buku populer tampil pada dashboard.", "Berhasil"],
        ],
        widths=[0.45, 1.65, 3.35, 0.8],
    )

    add_chapter(document, "V", "Penutup")
    add_heading(document, "Penutup", 1)
    add_heading(document, "Kesimpulan", 2)
    add_body(
        document,
        "Berdasarkan hasil analisis, perancangan, dan implementasi, aplikasi SHELFY berhasil dibuat sebagai aplikasi "
        "perpustakaan digital berbasis web dengan database NoSQL MongoDB. Aplikasi mendukung autentikasi, pembagian akses "
        "admin, pustakawan, dan mahasiswa, pengelolaan buku, pencarian dan filter, katalog, keranjang, checkout, peminjaman, "
        "pengembalian, perpanjangan, pembayaran denda, rekapitulasi, serta nota yang dapat dicetak atau disimpan sebagai PDF.",
    )
    add_body(
        document,
        "Penerapan Laravel MVC membuat source code lebih terstruktur karena logika dipisahkan ke Controller, data MongoDB "
        "dikelola melalui Model, dan tampilan disusun menggunakan Blade. Collection users, members, books, loans, dan cart_items "
        "mendukung proses utama aplikasi dengan referencing antar dokumen dan embedded document untuk riwayat perpanjangan.",
    )
    add_body(
        document,
        "Fitur rekapitulasi telah memenuhi kebutuhan Basis Data II melalui penggunaan operasi count, sum, average, greater "
        "than, filter, dan pengelompokan kategori. Dengan demikian, SHELFY dapat membantu proses perpustakaan menjadi lebih "
        "terpantau, sekaligus menunjukkan implementasi database MongoDB pada aplikasi web nyata.",
    )

    add_heading(document, "Saran", 2)
    add_body(
        document,
        "Pengembangan berikutnya dapat menambahkan notifikasi jatuh tempo melalui email atau WhatsApp, sistem reservasi ketika "
        "stok buku habis, dan audit log yang merekam setiap perubahan oleh petugas. Keamanan dapat ditingkatkan dengan "
        "verifikasi email, pengaturan hak akses yang lebih rinci, serta validasi file upload yang lebih lengkap.",
    )
    add_body(
        document,
        "Selain itu, nota PDF dapat dikembangkan menggunakan template resmi perpustakaan, katalog dapat dilengkapi rekomendasi "
        "buku berdasarkan riwayat peminjaman, dan dashboard dapat ditambah grafik interaktif untuk membantu analisis koleksi "
        "serta kebiasaan peminjaman mahasiswa.",
    )

    add_chapter(document, "", "Lampiran")
    # Remove the blank roman label created for the appendix.
    body = document._element.body
    appendix_title = None
    for paragraph in reversed(document.paragraphs):
        if paragraph.text.strip().upper() == "LAMPIRAN":
            appendix_title = paragraph
            break
    if appendix_title is not None:
        previous = appendix_title._p.getprevious()
        if previous is not None and element_text(previous).strip().upper() == "BAB":
            body.remove(previous)

    add_heading(document, "Daftar Path Screenshot Visual Studio Code", 1)
    add_body(
        document,
        "Gunakan daftar berikut ketika mengambil screenshot Visual Studio Code. Buka file pada path lengkap, tampilkan bagian "
        "method yang disebutkan, sembunyikan data sensitif pada .env, lalu tempel screenshot pada placeholder di BAB IV.",
    )
    add_table(
        document,
        ["No.", "Bagian", "Path Folder/File yang Dibuka"],
        [
            ["1", "Login dan register", ROOT / "app/Http/Controllers/Auth"],
            ["2", "Pengalihan role", ROOT / "app/Support/Shelfy.php"],
            ["3", "CRUD buku dan cover", ROOT / "app/Http/Controllers/BookController.php"],
            ["4", "Keranjang dan checkout", ROOT / "app/Http/Controllers/StudentController.php"],
            ["5", "Bukti buku diambil", ROOT / "app/Http/Controllers/LoanController.php"],
            ["6", "Pengembalian dan perpanjangan", ROOT / "app/Http/Controllers/ReturnController.php"],
            ["7", "Nota dan pembayaran", ROOT / "app/Http/Controllers/ReceiptController.php"],
            ["8", "Rekapitulasi", ROOT / "app/Http/Controllers/DashboardController.php"],
            ["9", "Koneksi MongoDB", ROOT / "config/database.php"],
            ["10", "Model/collection MongoDB", ROOT / "app/Models"],
            ["11", "Route aplikasi", ROOT / "routes/web.php"],
            ["12", "Struktur seluruh project", ROOT],
        ],
        widths=[0.45, 1.75, 4.05],
    )

    add_heading(document, "Struktur Folder Utama", 1)
    add_code_placeholder(
        document,
        "Struktur folder project SHELFY pada panel Explorer Visual Studio Code",
        [Path(".")],
        "Buka folder project SHELFY di Visual Studio Code. Expand folder app/Http/Controllers, app/Models, resources/views, routes, database, public, dan storage agar struktur project terlihat jelas.",
    )

    document.add_page_break()
    add_heading(document, "Ringkasan Source Code Penting", 1)
    add_table(
        document,
        ["Bagian", "File Utama", "Keterangan"],
        [
            ["Autentikasi", "app/Http/Controllers/Auth", "Login, register, reset password, dan session."],
            ["Buku", "BookController.php; Book.php", "CRUD, filter, stok, detail, dan cover."],
            ["Anggota", "MemberController.php; Member.php", "Direktori anggota hasil registrasi."],
            ["Peminjaman", "StudentController.php; LoanController.php", "Keranjang, checkout, filter, dan bukti pengambilan."],
            ["Pengembalian", "ReturnController.php", "Tanggal kembali manual, denda, dan perpanjangan."],
            ["Nota", "ReceiptController.php; receipts/show.blade.php", "Nota, pembayaran, konfirmasi, dan cetak PDF."],
            ["Rekap", "DashboardController.php; RecapController.php", "SUM, AVG, greater than, kategori, dan buku populer."],
            ["Route", "routes/web.php; routes/auth.php", "Daftar endpoint dan middleware role."],
            ["Database", "config/database.php; app/Models", "Koneksi MongoDB dan definisi collection."],
        ],
        widths=[1.15, 2.35, 2.75],
    )
    add_body(document, "Link GitHub: -")

    for table in document.tables:
        for index, row in enumerate(table.rows):
            prevent_row_split(row)
            if index == 0 and len(table.rows) > 2:
                set_repeat_header(row)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
