from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\Program Files\xampp\htdocs\laravel\shelfy")
OUTPUT = ROOT / "docs" / "Pembagian Tugas dan Panduan Presentasi SHELFY.docx"

NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
MUTED = RGBColor(85, 85, 85)
TEAL = RGBColor(0, 126, 121)


def set_cellless_page_border(section):
    sect_pr = section._sectPr
    pg_borders = sect_pr.find(qn("w:pgBorders"))
    if pg_borders is None:
        pg_borders = OxmlElement("w:pgBorders")
        pg_borders.set(qn("w:offsetFrom"), "page")
        sect_pr.append(pg_borders)

    for edge in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "24")
        element.set(qn("w:color"), "D9E2F3")
        pg_borders.append(element)


def configure_document(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = NAVY
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 14, 6),
        ("Heading 2", 13, BLUE, 12, 5),
        ("Heading 3", 11.5, TEAL, 9, 4),
    ]:
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ["List Bullet", "List Number"]:
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(10.5)
        style.font.color.rgb = NAVY
        style.paragraph_format.left_indent = Inches(0.32)
        style.paragraph_format.first_line_indent = Inches(-0.18)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.12

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("SHELFY | Panduan Presentasi TUBES Basis Data II")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("D4SIKC-48-01 | Panduan Presentasi SHELFY")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED


def add_title(document, text, size, color=NAVY, bold=True, after=6):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(after)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    return paragraph


def add_body(document, text, bold_label=None):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    if bold_label and text.startswith(bold_label):
        label = paragraph.add_run(bold_label)
        label.bold = True
        label.font.name = "Arial"
        label.font.size = Pt(10.5)
        label.font.color.rgb = NAVY
        rest = paragraph.add_run(text[len(bold_label):])
        rest.font.name = "Arial"
        rest.font.size = Pt(10.5)
        rest.font.color.rgb = NAVY
    else:
        run = paragraph.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(10.5)
        run.font.color.rgb = NAVY
    return paragraph


def add_bullet(document, text, level=0):
    paragraph = document.add_paragraph(text, style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.32 + 0.25 * level)
    paragraph.paragraph_format.first_line_indent = Inches(-0.18)
    return paragraph


def add_number(document, text):
    return document.add_paragraph(text, style="List Number")


def add_callout(document, label, text):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.left_indent = Inches(0.2)
    paragraph.paragraph_format.right_indent = Inches(0.2)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(7)

    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "EAF6F5")
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "007E79")
    borders.append(left)
    p_pr.append(borders)

    run = paragraph.add_run(f"{label}: ")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(10.5)
    run.font.color.rgb = TEAL
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10.5)
    run.font.color.rgb = NAVY
    return paragraph


def add_question(document, question, answer):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(f"Pertanyaan: {question}")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(10.5)
    run.font.color.rgb = BLUE

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(f"Jawaban yang dapat disampaikan: {answer}")
    run.font.name = "Arial"
    run.font.size = Pt(10.25)
    run.font.color.rgb = NAVY


def add_member_section(document, name, role, responsibilities, presentation, demo, files, questions, transition):
    document.add_page_break()
    document.add_heading(name, level=1)
    add_callout(document, "Peran utama", role)

    document.add_heading("Pembagian Tugas", level=2)
    for item in responsibilities:
        add_bullet(document, item)

    document.add_heading("Penjelasan yang Disampaikan Saat Presentasi", level=2)
    for item in presentation:
        add_body(document, item)

    document.add_heading("Urutan Demo", level=2)
    for item in demo:
        add_bullet(document, item)

    document.add_heading("File atau Bagian Kode yang Perlu Dibuka", level=2)
    for item in files:
        add_bullet(document, str(ROOT / item))

    document.add_heading("Kemungkinan Pertanyaan Dosen", level=2)
    for question, answer in questions:
        add_question(document, question, answer)

    add_callout(document, "Kalimat transisi", transition)


def build():
    document = Document()
    configure_document(document)

    add_title(document, "PEMBAGIAN TUGAS DAN PANDUAN PRESENTASI", 20, NAVY, True, 3)
    add_title(document, "APLIKASI PERPUSTAKAAN DIGITAL SHELFY", 17, TEAL, True, 12)
    add_title(document, "Tugas Besar Sistem Basis Data II", 12, MUTED, False, 4)
    add_title(document, "D4SIKC-48-01", 11, MUTED, True, 24)

    document.add_heading("Tujuan Dokumen", level=1)
    add_body(
        document,
        "Dokumen ini menjadi panduan pembagian tanggung jawab, urutan presentasi, demonstrasi aplikasi, dan jawaban singkat "
        "ketika dosen memberikan pertanyaan. Penyampaian tidak perlu dihafalkan kata per kata. Setiap anggota cukup memahami "
        "alur, alasan teknis, serta file source code yang menjadi tanggung jawabnya.",
    )

    document.add_heading("Pembagian Besar Presentasi", level=1)
    for item in [
        "Michael membuka presentasi, menjelaskan tujuan aplikasi, arsitektur, autentikasi, role, profil, serta integrasi akhir.",
        "Anantha menjelaskan katalog dan pengelolaan buku, MongoDB/Compass, seeder, dashboard, rekapitulasi, keranjang, dan checkout.",
        "Fadhil menjelaskan seluruh siklus transaksi peminjaman sampai pengembalian, denda, pembayaran, dan nota PDF.",
        "Mumpuni menjelaskan konsep NoSQL, desain collection, fitur katalog/filter dari sisi mahasiswa, testing, dan dokumentasi.",
    ]:
        add_number(document, item)

    document.add_heading("Catatan Pembagian Modul Buku", level=1)
    add_callout(
        document,
        "Jawaban jika ditanya mengenai pembagian yang beririsan",
        "Anantha menjadi penanggung jawab utama implementasi dan demonstrasi CRUD Book. Michael menangani integrasi akhir "
        "modul Book dengan autentikasi, role petugas, profil, storage, dan alur aplikasi secara keseluruhan. Pembagian ini "
        "merupakan kolaborasi integrasi, bukan duplikasi pekerjaan.",
    )

    document.add_heading("Aturan Menjawab Pertanyaan Dosen", level=1)
    for item in [
        "Jawab berdasarkan aplikasi yang benar-benar berjalan dan tunjukkan source code jika diminta.",
        "Mulai jawaban dari tujuan bisnis, lanjutkan alur proses, lalu jelaskan collection atau Controller yang terlibat.",
        "Jika pertanyaan berada di luar bagian sendiri, jawab bagian yang dipahami kemudian oper kepada penanggung jawab utama.",
        "Hindari menjawab hanya dengan nama fitur. Jelaskan input, proses backend, data MongoDB yang berubah, dan hasil pada UI.",
    ]:
        add_bullet(document, item)

    add_member_section(
        document,
        "Michael Eluzai Situmorang",
        "Koordinator; autentikasi Laravel Breeze; model User dan Member; role middleware; profil; integrasi akhir modul Book, cover, storage, dan seluruh aplikasi.",
        [
            "Mengkoordinasikan integrasi modul agar login, role, buku, anggota, transaksi, profil, dan storage dapat digunakan dalam satu aplikasi.",
            "Mengimplementasikan serta menjelaskan autentikasi Laravel Breeze untuk login, register, logout, session, dan pengalihan halaman berdasarkan role.",
            "Mengelola model User dan Member serta hubungan referencing antara akun login mahasiswa dan data keanggotaan perpustakaan.",
            "Mengimplementasikan middleware role untuk membatasi akses halaman admin/pustakawan dan mahasiswa.",
            "Mengelola halaman profil termasuk perubahan data diri dan upload foto profil ke storage Laravel.",
            "Melakukan integrasi akhir CRUD Book, cover buku, symbolic link storage, dan pengujian alur antarmodul.",
        ],
        [
            "SHELFY dibuat untuk mengatasi pengelolaan perpustakaan yang sebelumnya sulit dipantau secara manual. Aplikasi menangani katalog, anggota, peminjaman, pengembalian, perpanjangan, denda, dan nota dalam satu sistem.",
            "Arsitektur aplikasi menggunakan Laravel MVC. Route menerima permintaan, Controller menangani validasi dan proses bisnis, Model membaca atau mengubah dokumen MongoDB, lalu Blade menampilkan hasil kepada pengguna.",
            "Laravel Breeze digunakan sebagai kerangka autentikasi. Setelah login berhasil, Shelfy::homeRouteName menentukan apakah pengguna diarahkan ke dashboard petugas atau dashboard mahasiswa.",
            "Collection users menyimpan kebutuhan autentikasi dan otorisasi seperti email, password terenkripsi, role, status, serta member_id. Collection members menyimpan profil keanggotaan seperti NIM, prodi, kontak, alamat, dan status layanan.",
            "User dan Member dipisahkan agar data keamanan akun tidak tercampur dengan data operasional anggota. Keduanya dihubungkan menggunakan member_id dan user_id.",
            "Foto profil dan cover tidak disimpan sebagai binary di MongoDB. File fisik berada pada storage/app/public, sedangkan MongoDB hanya menyimpan path file. Perintah php artisan storage:link membuat file tersebut dapat diakses melalui public/storage.",
        ],
        [
            "Buka halaman login dan tunjukkan bahwa pengguna diarahkan sesuai role.",
            "Buka register mahasiswa, jelaskan bahwa register membuat dokumen users dan members lalu kembali ke login.",
            "Login sebagai mahasiswa dan buka Profil untuk menunjukkan perubahan data serta upload foto.",
            "Login sebagai admin/pustakawan dan tunjukkan pembatasan menu berdasarkan role.",
            "Buka struktur MVC dan tunjukkan Controller, Model, Blade view, routes, serta middleware.",
        ],
        [
            Path("app/Http/Controllers/Auth/AuthenticatedSessionController.php"),
            Path("app/Http/Controllers/Auth/RegisteredUserController.php"),
            Path("app/Http/Controllers/ProfileController.php"),
            Path("app/Models/User.php"),
            Path("app/Models/Member.php"),
            Path("app/Http/Middleware/EnsureAdmin.php"),
            Path("app/Http/Middleware/EnsureStudent.php"),
            Path("app/Support/Shelfy.php"),
            Path("routes/web.php"),
        ],
        [
            (
                "Mengapa users dan members dibuat menjadi dua collection?",
                "Users berfokus pada akun login, password, role, dan status akses. Members berfokus pada identitas serta layanan anggota perpustakaan. Pemisahan ini mengurangi pencampuran data sensitif dengan data operasional, tetapi keduanya tetap terhubung melalui referencing.",
            ),
            (
                "Bagaimana password disimpan?",
                "Password tidak disimpan dalam teks biasa. Laravel melakukan hashing melalui cast hashed atau Hash::make, kemudian login memverifikasi hash tersebut melalui mekanisme autentikasi Breeze.",
            ),
            (
                "Bagaimana pembatasan role bekerja?",
                "Route petugas menggunakan middleware admin yang memeriksa isStaff, sedangkan route mahasiswa menggunakan middleware student yang memeriksa isStudent. Jika role tidak sesuai, sistem mengembalikan HTTP 403.",
            ),
            (
                "Mengapa foto tidak disimpan langsung di MongoDB?",
                "Menyimpan file besar sebagai binary akan memperbesar dokumen dan menyulitkan akses publik. Karena itu file disimpan di storage Laravel, sedangkan MongoDB menyimpan avatar_path atau cover_path.",
            ),
            (
                "Apa kontribusi Michael pada modul Book jika Anantha menangani CRUD Book?",
                "Anantha menjadi implementer utama dan demonstrator CRUD Book. Michael mengintegrasikan modul tersebut dengan role petugas, storage, profil, navigasi, serta memastikan seluruh aplikasi bekerja bersama melalui pengujian akhir.",
            ),
        ],
        "Setelah arsitektur, autentikasi, dan role dijelaskan, bagian pengelolaan koleksi buku serta database MongoDB akan dilanjutkan oleh Anantha.",
    )

    add_member_section(
        document,
        "Muhammad Anantha Mahardika Ridwan",
        "Penanggung jawab utama Model/CRUD Book, cover dan storage link, konfigurasi MongoDB/Compass, seeder, dashboard, rekap agregasi, serta demo katalog, filter, detail, keranjang, dan checkout.",
        [
            "Mengimplementasikan Model Book dan operasi tambah, baca, ubah, serta hapus data buku.",
            "Mengimplementasikan pencarian judul/penulis/ISBN, filter kategori, filter status stok, dan halaman detail buku.",
            "Mengelola upload cover buku ke storage serta memastikan cover dapat ditampilkan melalui public/storage.",
            "Mengonfigurasi koneksi MongoDB lokal dan menjelaskan pemeriksaan database melalui MongoDB Compass.",
            "Mengelola seeder untuk data awal admin, pustakawan, mahasiswa, buku, anggota, dan contoh peminjaman.",
            "Mengimplementasikan dashboard dan rekap agregasi seperti jumlah, total stok, rata-rata, greater than, kategori, serta buku populer.",
            "Mendemonstrasikan katalog mahasiswa, masuk keranjang, dan checkout.",
        ],
        [
            "Model Book terhubung ke collection books. Field pentingnya meliputi judul, penulis, kategori, ISBN, stok_total, stok_tersedia, dipinjam_count, deskripsi, dan cover_path.",
            "BookController::store digunakan untuk tambah dan edit. Ketika stok total diedit, stok tersedia dihitung berdasarkan jumlah eksemplar yang sedang dipinjam agar data inventaris tetap konsisten.",
            "Filter dilakukan terhadap data buku berdasarkan kata kunci, kategori, dan status stok. Status tersedia atau habis dihitung dari stok_tersedia.",
            "MongoDB dihubungkan melalui package mongodb/laravel-mongodb. Konfigurasi menggunakan DB_CONNECTION=mongodb, host 127.0.0.1, port 27017, dan database shelfy_db.",
            "Seeder menyediakan data demo yang konsisten agar aplikasi dapat langsung dipresentasikan. Command php artisan shelfy:fresh membersihkan collection demo lalu mengisinya kembali.",
            "Dashboard melakukan rekapitulasi menggunakan count, sum, average, greater than, groupBy, dan pengurutan berdasarkan dipinjam_count untuk buku populer.",
            "Keranjang menggunakan collection cart_items sebagai penyimpanan sementara. Saat checkout, item divalidasi lalu dibuat menjadi dokumen loans dengan status menunggu_diambil.",
        ],
        [
            "Buka MongoDB Compass dan tunjukkan database shelfy_db beserta collection utama.",
            "Buka halaman Buku sebagai pustakawan, tambah satu buku, upload cover, lalu edit stok.",
            "Gunakan pencarian dan filter kategori/status, kemudian buka detail buku.",
            "Buka dashboard dan jelaskan kartu rekap serta daftar buku per kategori.",
            "Login sebagai mahasiswa, cari buku, buka detail, masukkan ke keranjang, lalu checkout.",
            "Kembali ke Compass dan tunjukkan perubahan cart_items serta loans.",
        ],
        [
            Path("app/Models/Book.php"),
            Path("app/Models/CartItem.php"),
            Path("app/Http/Controllers/BookController.php"),
            Path("app/Http/Controllers/DashboardController.php"),
            Path("app/Http/Controllers/RecapController.php"),
            Path("app/Http/Controllers/StudentController.php"),
            Path("app/Support/ShelfySeeder.php"),
            Path("config/database.php"),
            Path(".env.example"),
            Path("routes/console.php"),
        ],
        [
            (
                "Apa bukti aplikasi menggunakan MongoDB, bukan MySQL?",
                "Model utama menggunakan MongoDB Laravel Model, konfigurasi default memakai driver mongodb, dan data dapat diperlihatkan langsung pada database shelfy_db di MongoDB Compass.",
            ),
            (
                "Apa perbedaan stok_total dan stok_tersedia?",
                "stok_total adalah seluruh eksemplar fisik yang dimiliki. stok_tersedia adalah eksemplar yang dapat dipinjam saat ini. Selisihnya menunjukkan jumlah yang sedang berada pada transaksi aktif.",
            ),
            (
                "Bagaimana filter buku dilakukan?",
                "BookController mengambil data dari collection books, kemudian filterBooks mencocokkan kata kunci pada judul, penulis, dan ISBN, dilanjutkan filter kategori serta status stok.",
            ),
            (
                "Di mana operasi SUM, AVG, dan Greater Than digunakan?",
                "SUM digunakan untuk total stok, AVG untuk rata-rata stok per judul, dan Greater Than untuk menghitung buku dengan stok_total lebih dari lima. Hasilnya ditampilkan pada dashboard atau rekap.",
            ),
            (
                "Mengapa menggunakan cart_items sebelum loans?",
                "cart_items menjadi area sementara agar mahasiswa dapat memeriksa pilihan sebelum checkout. Setelah checkout berhasil, data dipindahkan secara logis menjadi loans dan item keranjang dihapus.",
            ),
            (
                "Bagaimana mencegah checkout buku yang stoknya habis?",
                "Stok diperiksa ketika buku dimasukkan ke keranjang dan diperiksa kembali saat checkout. Pemeriksaan kedua diperlukan karena stok dapat berubah selama buku berada di keranjang.",
            ),
        ],
        "Setelah buku berhasil di-checkout dan menjadi peminjaman, alur operasional pengambilan hingga nota akan dijelaskan oleh Fadhil.",
    )

    add_member_section(
        document,
        "Muhammad Fadhil Athallah",
        "Penanggung jawab Loan, bukti pengambilan, perpanjangan, pengembalian manual, denda, pembayaran, dan nota PDF.",
        [
            "Mengimplementasikan model dan alur status Loan dari menunggu_diambil, dipinjam, terlambat, hingga dikembalikan.",
            "Mengimplementasikan bukti buku sudah diambil oleh mahasiswa melalui konfirmasi pustakawan.",
            "Mengimplementasikan perpanjangan tanggal jatuh tempo manual beserta alasan dan riwayat perpanjangan.",
            "Mengimplementasikan pengembalian dengan tanggal manual, perubahan stok, kalkulasi keterlambatan, dan denda.",
            "Mengimplementasikan pembayaran denda mahasiswa serta konfirmasi lunas oleh pustakawan.",
            "Mengimplementasikan nota pengembalian yang dapat dicetak atau disimpan menjadi PDF.",
        ],
        [
            "Setelah mahasiswa checkout, Loan dibuat dengan status menunggu_diambil. Status ini berarti transaksi sudah tercatat, tetapi buku fisik belum dikonfirmasi diserahkan.",
            "Pustakawan mencatat bukti pengambilan melalui LoanController::pickup. Sistem menyimpan tanggal_diambil, bukti_pengambilan, petugas_pengambilan, lalu mengubah status menjadi dipinjam.",
            "Perpanjangan hanya dapat dilakukan untuk peminjaman aktif yang belum terlambat. Tanggal baru harus melewati jatuh tempo lama, dan setiap perubahan disimpan dalam riwayat_perpanjangan.",
            "Tanggal pengembalian dimasukkan manual oleh pustakawan agar sesuai kejadian nyata. Sistem menghitung selisih dari tanggal jatuh tempo melalui Shelfy::lateFee.",
            "Jika terlambat, total denda dihitung dari hari_terlambat dikali denda_per_hari. Jika tidak terlambat, pembayaran otomatis berstatus lunas atau bebas denda.",
            "Mahasiswa dapat mencatat pembayaran melalui Transfer atau QRIS. Pustakawan kemudian mengonfirmasi pembayaran menjadi lunas.",
            "Nota mengambil data dari dokumen loan yang sudah dikembalikan. Halaman nota menggunakan CSS print sehingga dapat dicetak atau disimpan sebagai PDF melalui browser.",
        ],
        [
            "Tunjukkan transaksi berstatus menunggu_diambil pada halaman Peminjaman.",
            "Klik konfirmasi bukti buku diambil dan tunjukkan status berubah menjadi dipinjam.",
            "Buka halaman Pengembalian dan demonstrasikan popup Perpanjang dengan tanggal serta alasan manual.",
            "Proses pengembalian menggunakan tanggal manual, kemudian jelaskan perhitungan denda.",
            "Buka nota hasil pengembalian dan tunjukkan tombol Cetak/Simpan PDF.",
            "Jika terdapat denda, login mahasiswa untuk memilih pembayaran lalu login pustakawan untuk konfirmasi lunas.",
        ],
        [
            Path("app/Models/Loan.php"),
            Path("app/Http/Controllers/LoanController.php"),
            Path("app/Http/Controllers/ReturnController.php"),
            Path("app/Http/Controllers/ReceiptController.php"),
            Path("app/Support/Shelfy.php"),
            Path("resources/views/shelfy/loans/index.blade.php"),
            Path("resources/views/shelfy/returns/index.blade.php"),
            Path("resources/views/shelfy/receipts/show.blade.php"),
        ],
        [
            (
                "Mengapa status awal setelah checkout adalah menunggu_diambil?",
                "Checkout mencatat permintaan peminjaman, tetapi belum membuktikan buku fisik sudah diserahkan. Status baru berubah menjadi dipinjam setelah pustakawan melakukan konfirmasi pengambilan.",
            ),
            (
                "Mengapa tanggal kembali harus diinput manual?",
                "Agar data sesuai kejadian nyata. Buku bisa diterima pada tanggal tertentu yang berbeda dari waktu aplikasi dibuka, sehingga pustakawan harus mencatat tanggal fisik pengembalian.",
            ),
            (
                "Bagaimana denda dihitung?",
                "Sistem menghitung selisih positif antara tanggal kembali dan tanggal jatuh tempo. Hari terlambat dikalikan denda per hari sebesar Rp2.000. Jika selisih tidak positif, denda nol.",
            ),
            (
                "Mengapa riwayat perpanjangan disimpan sebagai array di loans?",
                "Riwayat perpanjangan selalu menjadi bagian dari satu transaksi loan. Embedded array memudahkan melihat tanggal lama, tanggal baru, petugas, waktu proses, dan alasan tanpa collection tambahan.",
            ),
            (
                "Apakah nota benar-benar file PDF?",
                "Nota dibuat sebagai halaman HTML khusus cetak. Pengguna memilih Cetak/Simpan PDF, kemudian browser menghasilkan PDF. Cara ini tetap menghasilkan dokumen PDF formal tanpa menyimpan file PDF permanen untuk setiap transaksi.",
            ),
            (
                "Siapa yang boleh mengonfirmasi pembayaran?",
                "Mahasiswa hanya memilih dan mencatat metode pembayaran. Konfirmasi lunas dilakukan admin atau pustakawan agar status pembayaran memiliki verifikasi petugas.",
            ),
        ],
        "Setelah siklus transaksi selesai, Mumpuni akan menjelaskan alasan penggunaan NoSQL, desain collection, pencarian/filter, serta hasil testing aplikasi.",
    )

    add_member_section(
        document,
        "Mumpuni Nur Idzati Ayuningtyas",
        "Penanggung jawab katalog, pencarian/filter, keranjang dan checkout dari sisi pengguna, testing, dokumentasi, konsep NoSQL, desain collection, agregasi, serta seeder.",
        [
            "Menjelaskan dasar pemilihan MongoDB sebagai database NoSQL berbasis dokumen.",
            "Mendokumentasikan desain collection users, members, books, cart_items, dan loans beserta relasinya.",
            "Mengelola dan mendemonstrasikan katalog mahasiswa, pencarian/filter, keranjang, dan checkout dari sisi pengalaman pengguna.",
            "Menjelaskan implementasi rekapitulasi SUM, AVG, Greater Than, pengelompokan kategori, dan buku populer.",
            "Menyusun serta menjalankan pengujian alur aplikasi agar fitur utama dan pembagian role tidak mengalami regresi.",
            "Menyusun dokumentasi laporan, bukti screenshot, hasil pengujian, dan penjelasan seeder.",
        ],
        [
            "MongoDB adalah database NoSQL berbasis dokumen. Data disimpan dalam format BSON yang fleksibel sehingga setiap dokumen dapat memiliki field yang berkembang sesuai kebutuhan aplikasi.",
            "SHELFY menggunakan referencing untuk menghubungkan dokumen, misalnya users.member_id menuju members dan loans.book_id menuju books. Beberapa informasi seperti judul_buku dan nama_anggota juga diduplikasi secara terkontrol agar riwayat tetap mudah dibaca.",
            "Collection users menyimpan akun dan role; members menyimpan identitas anggota; books menyimpan katalog dan stok; cart_items menyimpan pilihan sementara; loans menyimpan seluruh siklus transaksi.",
            "Pencarian katalog membantu mahasiswa menemukan buku berdasarkan judul, penulis, atau ISBN. Filter kategori dan status stok mempersempit hasil tanpa mengubah data asli.",
            "Testing dilakukan menggunakan Pest/Laravel Feature Test. Pengujian mencakup autentikasi, role, profil, CRUD buku, anggota, checkout, pengambilan, perpanjangan, pengembalian, pembayaran, dan nota.",
            "Seeder digunakan agar data demo selalu dapat dipulihkan dengan struktur konsisten. Ini memudahkan presentasi dan pengujian tanpa harus memasukkan seluruh data secara manual.",
        ],
        [
            "Buka MongoDB Compass dan jelaskan isi serta tujuan setiap collection.",
            "Buka katalog mahasiswa dan demonstrasikan pencarian serta filter.",
            "Tunjukkan keranjang sebagai data sementara, lalu jelaskan perubahan menjadi loan ketika checkout.",
            "Buka dashboard atau rekap untuk menunjukkan SUM, AVG, Greater Than, kategori, dan buku populer.",
            "Tunjukkan hasil perintah php artisan test dan jelaskan skenario yang diuji.",
            "Jelaskan penggunaan php artisan shelfy:fresh untuk memulihkan data demo.",
        ],
        [
            Path("app/Models"),
            Path("app/Http/Controllers/StudentController.php"),
            Path("app/Http/Controllers/DashboardController.php"),
            Path("app/Http/Controllers/RecapController.php"),
            Path("app/Support/ShelfySeeder.php"),
            Path("tests/Feature/ShelfyFlowTest.php"),
            Path("tests/Feature/ProfileTest.php"),
            Path("tests/Feature/Auth"),
            Path("routes/console.php"),
            Path("docs"),
        ],
        [
            (
                "Mengapa memilih MongoDB untuk aplikasi perpustakaan?",
                "MongoDB sesuai karena struktur data transaksi dapat berkembang, misalnya penambahan bukti pengambilan, riwayat perpanjangan, pembayaran, dan nota. MongoDB juga memenuhi ketentuan tugas yang mewajibkan database NoSQL.",
            ),
            (
                "Apakah MongoDB tidak memiliki relasi?",
                "MongoDB tidak menggunakan foreign key seperti database relasional, tetapi tetap dapat menghubungkan data melalui referencing. SHELFY menyimpan id dokumen terkait dan melakukan denormalisasi pada informasi yang perlu dipertahankan dalam riwayat.",
            ),
            (
                "Apa contoh CRUD dalam aplikasi?",
                "Create saat menambah buku, register, dan checkout; Read saat menampilkan katalog dan anggota; Update saat mengedit buku, profil, status loan, atau pembayaran; Delete saat admin menghapus buku dan mahasiswa menghapus item keranjang.",
            ),
            (
                "Apa bukti fitur filter dan agregasi sudah diterapkan?",
                "Filter dapat didemonstrasikan langsung pada katalog atau halaman buku. Agregasi terlihat pada dashboard melalui total stok, rata-rata stok, buku dengan stok lebih dari lima, pengelompokan kategori, dan buku populer.",
            ),
            (
                "Bagaimana memastikan aplikasi tidak error setelah perubahan?",
                "Kami menjalankan automated feature test menggunakan php artisan test serta melakukan browser QA. Test memverifikasi status respons, hak akses, perubahan dokumen MongoDB, stok, status transaksi, dan file upload.",
            ),
            (
                "Apa fungsi seeder dan apakah aman digunakan?",
                "Seeder mengisi data demo awal yang konsisten. Command shelfy:fresh digunakan hanya saat ingin mereset data demo karena command tersebut menghapus collection aplikasi sebelum mengisi ulang.",
            ),
        ],
        "Demonstrasi dan penjelasan teknis telah selesai. Selanjutnya kelompok menyampaikan kesimpulan dan membuka sesi tanya jawab.",
    )

    document.add_page_break()
    document.add_heading("Pertanyaan Umum untuk Seluruh Kelompok", level=1)
    shared_questions = [
        (
            "Apa alur utama aplikasi SHELFY?",
            "Mahasiswa register dan login, mencari buku, memasukkan buku ke keranjang, lalu checkout. Pustakawan mengonfirmasi pengambilan, dapat memperpanjang, memproses pengembalian, dan mengonfirmasi pembayaran. Setelah selesai, mahasiswa memperoleh nota.",
        ),
        (
            "Collection mana yang paling kompleks?",
            "Collection loans paling kompleks karena menyimpan identitas transaksi, status, tanggal, bukti pengambilan, riwayat perpanjangan, data denda, pembayaran, dan nomor nota.",
        ),
        (
            "Bagaimana konsistensi stok dijaga?",
            "Stok diperiksa sebelum cart dan checkout, stok_tersedia berkurang setelah checkout, lalu bertambah ketika pengembalian diproses. Penghapusan buku juga ditolak jika masih memiliki loan aktif.",
        ),
        (
            "Apa kekurangan aplikasi saat ini?",
            "Belum terdapat notifikasi otomatis jatuh tempo, audit log lengkap, reservasi saat stok habis, dan integrasi pembayaran nyata. Fitur tersebut menjadi saran pengembangan berikutnya.",
        ),
        (
            "Apa hal yang menunjukkan kompleksitas database?",
            "Terdapat lima collection utama yang saling terhubung, role-based access, siklus status transaksi, referencing dan denormalisasi, embedded riwayat perpanjangan, perhitungan denda, serta agregasi dashboard.",
        ),
    ]
    for question, answer in shared_questions:
        add_question(document, question, answer)

    document.add_heading("Checklist Sebelum Presentasi", level=1)
    for item in [
        "Pastikan MongoDB Server aktif dan MongoDB Compass dapat membuka shelfy_db.",
        "Pastikan aplikasi berjalan di http://127.0.0.1:8000.",
        "Pastikan public/storage tersedia melalui php artisan storage:link.",
        "Siapkan akun admin, pustakawan, dan mahasiswa untuk demo.",
        "Pastikan data anggota, buku, dan transaksi demo tersedia.",
        "Jalankan php artisan test dan simpan hasil lulus untuk diperlihatkan.",
        "Buka file source code penting terlebih dahulu pada tab Visual Studio Code.",
        "Simpan salinan database atau gunakan seeder apabila data demo perlu dipulihkan.",
    ]:
        add_bullet(document, item)

    add_callout(
        document,
        "Penutup yang dapat disampaikan",
        "SHELFY telah menerapkan database NoSQL MongoDB, CRUD, filter, rekapitulasi, autentikasi, pembagian role, dan alur "
        "transaksi perpustakaan secara lengkap. Setiap anggota memahami bagian implementasi serta hubungan antarmodul yang "
        "telah dikerjakan bersama.",
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
